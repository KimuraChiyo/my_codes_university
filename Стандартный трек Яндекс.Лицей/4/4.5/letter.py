from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


document = Document()

# заголовок
style = document.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.font.italic = True
paragraph1 = document.add_paragraph('Dear Kate and Nick,')
p_fmt = paragraph1.paragraph_format
p_fmt.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

style = document.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
paragraph2 = document.add_paragraph('We are looking forward very much to your visit to our country this summer. We are expecting you at the beginning of July and are hoping that you may stay until the end of the month or longer.')







document.save('letter.docx')
